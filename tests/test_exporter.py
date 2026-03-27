import io

from docx import Document

from telegram_summarizer.exporter import export_docx, export_markdown, export_pdf, export_tlg


class TestExportMarkdown:
    def test_returns_text_as_is(self):
        text = "# Hello\n\nSome summary text."
        assert export_markdown(text) == text

    def test_empty_string(self):
        assert export_markdown("") == ""

    def test_preserves_unicode(self):
        text = "Привет мир! 你好世界 🌍"
        assert export_markdown(text) == text


class TestExportTlg:
    def test_plain_text_escapes_special_chars(self):
        assert export_tlg("hello.world!") == "hello\\.world\\!"

    def test_heading_converted_to_bold(self):
        assert export_tlg("# Title") == "*Title*"

    def test_heading_level2(self):
        assert export_tlg("## Subtitle") == "*Subtitle*"

    def test_heading_escapes_special_chars(self):
        assert export_tlg("# Hello.World!") == "*Hello\\.World\\!*"

    def test_bold_converted(self):
        assert export_tlg("some **bold** text") == "some *bold* text"

    def test_bold_with_special_chars(self):
        assert export_tlg("**hello!**") == "*hello\\!*"

    def test_bullet_list(self):
        result = export_tlg("- item one\n- item two")
        assert result == "• item one\n• item two"

    def test_bullet_list_with_special_chars(self):
        assert export_tlg("- hello.world") == "• hello\\.world"

    def test_empty_string(self):
        assert export_tlg("") == ""

    def test_multiline(self):
        text = "# Title\n\nSome text.\n\n- item"
        result = export_tlg(text)
        lines = result.split("\n")
        assert lines[0] == "*Title*"
        assert lines[1] == ""
        assert lines[2] == "Some text\\."
        assert lines[3] == ""
        assert lines[4] == "• item"

    def test_preserves_unicode(self):
        text = "Привет мир! 🌍"
        result = export_tlg(text)
        assert "Привет мир\\!" in result
        assert "🌍" in result

    def test_all_special_chars_escaped(self):
        for ch in "_*[]()~`>#+\\-=|{}.!":
            result = export_tlg(f"a{ch}b")
            assert f"a\\{ch}b" in result


class TestExportPdf:
    def test_returns_bytes(self):
        result = export_pdf("Hello world")
        assert isinstance(result, bytes)

    def test_starts_with_pdf_header(self):
        result = export_pdf("Hello world")
        assert result[:5] == b"%PDF-"

    def test_non_empty_output(self):
        result = export_pdf("Some text here")
        assert len(result) > 100

    def test_handles_headings(self):
        text = "# Title\n\n## Subtitle\n\nBody text."
        result = export_pdf(text)
        assert result[:5] == b"%PDF-"

    def test_handles_empty_string(self):
        result = export_pdf("")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    def test_handles_long_text(self):
        text = "Long line. " * 500
        result = export_pdf(text)
        assert isinstance(result, bytes)
        assert len(result) > 100

    def test_handles_multiline(self):
        text = "\n".join(f"Line {i}" for i in range(100))
        result = export_pdf(text)
        assert result[:5] == b"%PDF-"


class TestExportDocx:
    def test_returns_bytes(self):
        result = export_docx("Hello world")
        assert isinstance(result, bytes)

    def test_valid_docx(self):
        result = export_docx("Hello world")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Hello world"

    def test_handles_headings(self):
        text = "# Title\n\nBody text."
        result = export_docx(text)
        doc = Document(io.BytesIO(result))
        # Heading + paragraph
        assert len(doc.paragraphs) >= 2

    def test_handles_empty_string(self):
        result = export_docx("")
        assert isinstance(result, bytes)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 0

    def test_handles_unicode(self):
        text = "Привет мир! 你好世界"
        result = export_docx(text)
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].text == text

    def test_handles_long_text(self):
        text = "\n".join(f"Paragraph {i}" for i in range(100))
        result = export_docx(text)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 100

    def test_heading_levels(self):
        text = "# H1\n## H2\n### H3\nBody"
        result = export_docx(text)
        doc = Document(io.BytesIO(result))
        # 3 headings + 1 paragraph
        assert len(doc.paragraphs) == 4
