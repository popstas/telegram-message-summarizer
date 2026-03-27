import io
import os
import re

from docx import Document
from fpdf import FPDF

_DEJAVU_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
_DEJAVU_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


def export_markdown(summary_text: str) -> str:
    return summary_text


# Characters that must be escaped in Telegram MarkdownV2
_TLG_SPECIAL = r"_*[]()~`>#+\-=|{}.!"


def _escape_tlg(text: str) -> str:
    """Escape special characters for Telegram MarkdownV2."""
    result = []
    for ch in text:
        if ch in _TLG_SPECIAL:
            result.append("\\")
        result.append(ch)
    return "".join(result)


def _convert_inline(text: str) -> str:
    """Convert inline markdown (bold, italic) and escape the rest."""
    # Split on bold **text** first
    bold_parts = re.split(r"\*\*(.+?)\*\*", text)
    result = []
    for i, part in enumerate(bold_parts):
        if i % 2 == 1:
            # Bold content — convert to *text*
            result.append(f"*{_convert_italic(part)}*")
        else:
            result.append(_convert_italic(part))
    return "".join(result)


def _convert_italic(text: str) -> str:
    """Convert _italic_ markers and escape the rest."""
    parts = re.split(r"(?<!\w)_(.+?)_(?!\w)", text)
    result = []
    for i, part in enumerate(parts):
        if i % 2 == 1:
            # Italic content
            result.append(f"_{_escape_tlg(part)}_")
        else:
            result.append(_escape_tlg(part))
    return "".join(result)


def _convert_inline_code(text: str) -> str:
    """Handle inline `code` spans: preserve backticks, don't escape content inside."""
    parts = re.split(r"(`[^`]+`)", text)
    result = []
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            result.append(part)  # code span: pass through verbatim
        else:
            result.append(_convert_inline(part))
    return "".join(result)


def export_tlg(summary_text: str) -> str:
    """Convert standard markdown to Telegram MarkdownV2 format."""
    lines = summary_text.split("\n")
    out = []
    in_code_block = False
    for line in lines:
        # Handle fenced code blocks
        fence_match = re.match(r"^(\s*)```", line)
        if fence_match:
            if in_code_block:
                out.append(fence_match.group(1) + "```")
                in_code_block = False
            else:
                out.append(fence_match.group(1) + "```")
                in_code_block = True
            continue
        if in_code_block:
            out.append(line)
            continue

        # Convert headings: # Heading → *Heading*
        heading_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading_match:
            # Strip bold markers since the whole heading is already bold
            raw = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_match.group(2).strip())
            heading_text = _convert_inline_code(raw)
            out.append(f"*{heading_text}*")
            continue

        # Convert bold/italic and escape the rest
        # Check for bullet lists first
        bullet_match = re.match(r"^(\s*)-\s+(.*)", line)
        if bullet_match:
            indent = bullet_match.group(1)
            rest = _convert_inline_code(bullet_match.group(2))
            out.append(f"{indent}• {rest}")
        else:
            out.append(_convert_inline_code(line))

    return "\n".join(out)


def _setup_pdf_fonts(pdf: FPDF) -> str:
    """Set up fonts, return font family name to use."""
    if os.path.exists(_DEJAVU_REGULAR) and os.path.exists(_DEJAVU_BOLD):
        pdf.add_font("DejaVu", "", _DEJAVU_REGULAR)
        pdf.add_font("DejaVu", "B", _DEJAVU_BOLD)
        return "DejaVu"
    return "Helvetica"


def export_pdf(summary_text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    font_family = _setup_pdf_fonts(pdf)
    pdf.set_font(font_family, size=11)
    pdf.set_auto_page_break(auto=True, margin=15)

    w = pdf.epw
    for line in summary_text.split("\n"):
        if re.match(r"^#{1,3}\s", line):
            pdf.set_font(font_family, "B", size=14)
            pdf.multi_cell(w, 8, re.sub(r"^#{1,3}\s+", "", line).strip())
            pdf.set_font(font_family, size=11)
        elif line.strip() == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(w, 8, line)

    return bytes(pdf.output())


def export_docx(summary_text: str) -> bytes:
    doc = Document()

    for line in summary_text.split("\n"):
        if re.match(r"^#{1,3}\s", line):
            level = len(line) - len(line.lstrip("#"))
            doc.add_heading(re.sub(r"^#{1,3}\s+", "", line).strip(), level=min(level, 9))
        elif line.strip() == "":
            continue
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
